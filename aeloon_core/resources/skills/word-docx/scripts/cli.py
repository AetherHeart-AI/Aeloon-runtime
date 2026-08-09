#!/usr/bin/env python3
"""Create, edit, validate, and render DOCX files using Python."""

from __future__ import annotations

import argparse
import copy
import json
import os
import platform
import posixpath
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from collections.abc import Sequence
from datetime import UTC, datetime
from pathlib import Path, PurePosixPath
from typing import Any
from urllib.parse import urlparse

from lxml import etree

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
XML_NS = "http://www.w3.org/XML/1998/namespace"
NS = {"w": W_NS, "r": R_NS}
COMMENTS_REL_TYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
)
COMMENTS_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
)
SUPPORTED_SCHEMAS = {"word-edit-spec/v1", "word-edit-spec/v1.1"}
SCHEMA_OPERATIONS = {
    "word-edit-spec/v1": {"replace", "fill"},
    "word-edit-spec/v1.1": {"replace", "fill", "track_replace", "comment"},
}
SCOPE_PATTERNS = {
    "document": (re.compile(r"^word/document\.xml$"),),
    "headers": (re.compile(r"^word/header\d+\.xml$"),),
    "footers": (re.compile(r"^word/footer\d+\.xml$"),),
    "all": (
        re.compile(r"^word/document\.xml$"),
        re.compile(r"^word/header\d+\.xml$"),
        re.compile(r"^word/footer\d+\.xml$"),
    ),
}


class WordSkillError(ValueError):
    """A user-correctable Word skill failure."""


def _w(local: str) -> str:
    return f"{{{W_NS}}}{local}"


def _rel(local: str) -> str:
    return f"{{{REL_NS}}}{local}"


def _ct(local: str) -> str:
    return f"{{{CT_NS}}}{local}"


def _utc_now() -> str:
    return datetime.now(UTC).replace(microsecond=0).isoformat().replace("+00:00", "Z")


def _local_path(value: str, *, must_exist: bool = True) -> Path:
    parsed = urlparse(value)
    if parsed.scheme or parsed.netloc:
        raise WordSkillError("only local filesystem paths are supported; URI input is refused")
    path = Path(value).expanduser().resolve()
    if must_exist and not path.is_file():
        raise WordSkillError(f"input file does not exist: {path}")
    return path


def _input_docx(value: str) -> Path:
    path = _local_path(value)
    if path.suffix.lower() != ".docx":
        if path.suffix.lower() in {".doc", ".wps"}:
            raise WordSkillError(
                "native .doc/.wps is unsupported; use Word or WPS to save as .docx first"
            )
        raise WordSkillError("input must be an OOXML .docx file")
    return path


def _output_docx(
    value: str, *, source: Path | None = None, overwrite: bool = False
) -> Path:
    path = _local_path(value, must_exist=False)
    if path.suffix.lower() != ".docx":
        raise WordSkillError("output must use the .docx extension")
    if source is not None and path == source:
        raise WordSkillError("input and output paths must be different")
    if path.exists() and not overwrite:
        raise WordSkillError(f"output already exists (pass --overwrite to replace it): {path}")
    path.parent.mkdir(parents=True, exist_ok=True)
    return path


def _atomic_replace(temp_path: Path, output: Path) -> None:
    os.replace(temp_path, output)


def _write_docx_atomic(document: Any, output: Path) -> None:
    handle, temp_name = tempfile.mkstemp(
        prefix=f".{output.stem}-", suffix=".docx", dir=output.parent
    )
    os.close(handle)
    temp_path = Path(temp_name)
    try:
        document.save(temp_path)
        _atomic_replace(temp_path, output)
    finally:
        temp_path.unlink(missing_ok=True)


def _add_field(paragraph: Any, instruction: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Update field in Word"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, text, end))


def _add_hyperlink(paragraph: Any, label: str, url: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    style = OxmlElement("w:rStyle")
    style.set(qn("w:val"), "Hyperlink")
    properties.append(style)
    text = OxmlElement("w:t")
    text.text = label
    run.extend((properties, text))
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_inline(paragraph: Any, token: Any, source_dir: Path) -> None:
    children = token.children or []
    bold = False
    italic = False
    strike = False
    link_href: str | None = None
    link_text: list[str] = []

    for child in children:
        kind = child.type
        if kind == "strong_open":
            bold = True
        elif kind == "strong_close":
            bold = False
        elif kind == "em_open":
            italic = True
        elif kind == "em_close":
            italic = False
        elif kind == "s_open":
            strike = True
        elif kind == "s_close":
            strike = False
        elif kind == "link_open":
            link_href = child.attrGet("href")
            link_text = []
        elif kind == "link_close":
            if link_href is not None:
                _add_hyperlink(paragraph, "".join(link_text), link_href)
            link_href = None
            link_text = []
        elif kind == "image":
            source = child.attrGet("src") or ""
            parsed = urlparse(source)
            if parsed.scheme or parsed.netloc:
                raise WordSkillError(f"remote Markdown image is refused: {source}")
            image_path = (source_dir / source).resolve()
            if not image_path.is_file():
                raise WordSkillError(f"Markdown image does not exist: {image_path}")
            run = paragraph.add_run()
            run.add_picture(str(image_path))
        elif kind in {"softbreak", "hardbreak"}:
            if link_href is not None:
                link_text.append(" ")
            else:
                paragraph.add_run().add_break()
        elif kind == "code_inline":
            if link_href is not None:
                link_text.append(child.content)
            else:
                run = paragraph.add_run(child.content)
                run.font.name = "Consolas"
        elif kind == "text":
            if link_href is not None:
                link_text.append(child.content)
            else:
                run = paragraph.add_run(child.content)
                run.bold = bold
                run.italic = italic
                run.font.strike = strike
        elif kind == "html_inline" and child.content.lower() in {"<br>", "<br/>"}:
            paragraph.add_run().add_break()


def _table_rows(tokens: Sequence[Any], start: int) -> tuple[list[list[Any]], int]:
    rows: list[list[Any]] = []
    row: list[Any] | None = None
    cell_inline: Any | None = None
    index = start
    while index < len(tokens):
        token = tokens[index]
        if token.type == "table_close":
            return rows, index
        if token.type == "tr_open":
            row = []
        elif token.type == "inline" and row is not None:
            cell_inline = token
        elif token.type in {"th_close", "td_close"} and row is not None:
            row.append(cell_inline)
            cell_inline = None
        elif token.type == "tr_close" and row is not None:
            rows.append(row)
            row = None
        index += 1
    raise WordSkillError("Markdown table was not closed")


def _markdown_document(source: Path, args: argparse.Namespace) -> Any:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt
    from markdown_it import MarkdownIt

    markdown = source.read_text(encoding="utf-8")
    parser = MarkdownIt("commonmark").enable("table").enable("strikethrough")
    tokens = parser.parse(markdown)
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    normal = document.styles["Normal"]
    normal.font.name = args.font
    normal.font.size = Pt(args.font_size)

    if args.header:
        section.header.paragraphs[0].text = args.header
    footer = section.footer.paragraphs[0]
    if args.footer:
        footer.add_run(args.footer)
    if args.page_numbers:
        if args.footer:
            footer.add_run("  ·  ")
        _add_field(footer, "PAGE")

    if args.title:
        title = document.add_paragraph(style="Title")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.add_run(args.title)
        if args.subtitle:
            subtitle = document.add_paragraph()
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle.add_run(args.subtitle)
        if args.author:
            author = document.add_paragraph()
            author.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author.add_run(args.author)
        if args.date:
            date = document.add_paragraph()
            date.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date.add_run(args.date)
        document.add_page_break()

    if args.toc:
        document.add_heading("目录", level=1)
        _add_field(document.add_paragraph(), 'TOC \\o "1-3" \\h \\z \\u')
        document.add_page_break()

    list_stack: list[str] = []
    blockquote_depth = 0
    index = 0
    while index < len(tokens):
        token = tokens[index]
        kind = token.type
        if kind == "heading_open":
            level = min(int(token.tag[1:]), 9)
            paragraph = document.add_heading(level=level)
            _add_inline(paragraph, tokens[index + 1], source.parent)
            index += 2
        elif kind == "paragraph_open":
            style = None
            if list_stack:
                style = "List Number" if list_stack[-1] == "ordered" else "List Bullet"
            paragraph = document.add_paragraph(style=style)
            if blockquote_depth:
                paragraph.style = "Quote"
            _add_inline(paragraph, tokens[index + 1], source.parent)
            index += 2
        elif kind == "bullet_list_open":
            list_stack.append("bullet")
        elif kind == "ordered_list_open":
            list_stack.append("ordered")
        elif kind in {"bullet_list_close", "ordered_list_close"}:
            if list_stack:
                list_stack.pop()
        elif kind == "blockquote_open":
            blockquote_depth += 1
        elif kind == "blockquote_close":
            blockquote_depth = max(0, blockquote_depth - 1)
        elif kind in {"fence", "code_block"}:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(token.content.rstrip("\n"))
            run.font.name = "Consolas"
        elif kind == "hr":
            document.add_paragraph("―" * 24)
        elif kind in {"html_block", "html_inline"} and "pagebreak" in token.content.lower():
            document.add_page_break()
        elif kind == "table_open":
            rows, end_index = _table_rows(tokens, index + 1)
            if rows:
                width = max(len(row) for row in rows)
                table = document.add_table(rows=len(rows), cols=width)
                table.style = "Table Grid"
                for row_index, row in enumerate(rows):
                    for column_index, inline in enumerate(row):
                        cell = table.cell(row_index, column_index)
                        cell.text = ""
                        if inline is not None:
                            _add_inline(cell.paragraphs[0], inline, source.parent)
            index = end_index
        index += 1

    if len(document.sections) > 1:
        for extra_section in document.sections[1:]:
            extra_section.start_type = WD_SECTION.NEW_PAGE
    return document


def build_document(args: argparse.Namespace) -> dict[str, Any]:
    source = _local_path(args.input)
    if source.suffix.lower() not in {".md", ".markdown", ".txt"}:
        raise WordSkillError("build input must be Markdown or UTF-8 text")
    output = _output_docx(args.output, source=source, overwrite=args.overwrite)
    document = _markdown_document(source, args)
    _write_docx_atomic(document, output)
    return {
        "status": "created",
        "output": str(output),
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
    }


def _load_package(path: Path) -> dict[str, bytes]:
    try:
        with zipfile.ZipFile(path) as archive:
            bad_member = archive.testzip()
            if bad_member:
                raise WordSkillError(f"corrupt ZIP member: {bad_member}")
            return {name: archive.read(name) for name in archive.namelist()}
    except zipfile.BadZipFile as exc:
        raise WordSkillError(f"invalid DOCX ZIP package: {path}") from exc


def _xml(data: bytes, part_name: str) -> etree._Element:
    parser = etree.XMLParser(resolve_entities=False, no_network=True)
    try:
        return etree.fromstring(data, parser)
    except etree.XMLSyntaxError as exc:
        raise WordSkillError(f"invalid XML in {part_name}: {exc}") from exc


def _serialize(root: etree._Element) -> bytes:
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def _scoped_parts(entries: dict[str, bytes], scope: str) -> list[str]:
    patterns = SCOPE_PATTERNS.get(scope)
    if patterns is None:
        raise WordSkillError(f"unknown scope {scope!r}; use document, headers, footers, or all")
    return [name for name in entries if any(pattern.match(name) for pattern in patterns)]


def _set_text(node: etree._Element, value: str) -> None:
    node.text = value
    space = f"{{{XML_NS}}}space"
    if value[:1].isspace() or value[-1:].isspace():
        node.set(space, "preserve")
    else:
        node.attrib.pop(space, None)


def _matches(text: str, find: str, case_sensitive: bool, replace_all: bool) -> list[int]:
    haystack = text if case_sensitive else text.lower()
    needle = find if case_sensitive else find.lower()
    starts: list[int] = []
    offset = 0
    while True:
        found = haystack.find(needle, offset)
        if found < 0:
            break
        starts.append(found)
        if not replace_all:
            break
        offset = found + len(needle)
    return starts


def _replace_in_paragraph(
    paragraph: etree._Element,
    find: str,
    replacement: str,
    *,
    case_sensitive: bool,
    replace_all: bool,
) -> int:
    nodes = paragraph.xpath(".//w:t[not(ancestor::w:del)]", namespaces=NS)
    if not nodes:
        return 0
    values = [node.text or "" for node in nodes]
    full_text = "".join(values)
    starts = _matches(full_text, find, case_sensitive, replace_all)
    if not starts:
        return 0

    offsets: list[int] = []
    cursor = 0
    for value in values:
        offsets.append(cursor)
        cursor += len(value)

    for start in reversed(starts):
        end = start + len(find)
        first = next(
            index
            for index, offset in enumerate(offsets)
            if offset + len(values[index]) > start
        )
        last = max(
            index
            for index, offset in enumerate(offsets)
            if offset < end and offset + len(values[index]) > start
        )
        first_local = start - offsets[first]
        last_local = end - offsets[last]
        if first == last:
            updated = (
                (nodes[first].text or "")[:first_local]
                + replacement
                + (nodes[first].text or "")[last_local:]
            )
            _set_text(nodes[first], updated)
        else:
            prefix = (nodes[first].text or "")[:first_local]
            suffix = (nodes[last].text or "")[last_local:]
            _set_text(nodes[first], prefix + replacement)
            for index in range(first + 1, last):
                _set_text(nodes[index], "")
            _set_text(nodes[last], suffix)
    return len(starts)


def _direct_run_map(
    paragraph: etree._Element,
) -> tuple[list[etree._Element], list[str], list[int]]:
    runs = [child for child in paragraph if child.tag == _w("r")]
    values = ["".join(run.xpath(".//w:t/text()", namespaces=NS)) for run in runs]
    offsets: list[int] = []
    cursor = 0
    for value in values:
        offsets.append(cursor)
        cursor += len(value)
    return runs, values, offsets


def _text_run(template: etree._Element, text: str, *, deleted: bool = False) -> etree._Element:
    run = etree.Element(_w("r"))
    properties = template.find(_w("rPr"))
    if properties is not None:
        run.append(copy.deepcopy(properties))
    text_node = etree.SubElement(run, _w("delText") if deleted else _w("t"))
    _set_text(text_node, text)
    return run


def _affected_runs(
    paragraph: etree._Element, start: int, end: int
) -> tuple[list[etree._Element], list[str], list[int], int, int]:
    runs, values, offsets = _direct_run_map(paragraph)
    candidates = [
        index
        for index, offset in enumerate(offsets)
        if offset < end and offset + len(values[index]) > start
    ]
    if not candidates:
        raise WordSkillError(
            "target text is inside an unsupported nested run (for example a hyperlink)"
        )
    first, last = candidates[0], candidates[-1]
    if offsets[first] > start or offsets[last] + len(values[last]) < end:
        raise WordSkillError("target crosses an unsupported non-run OOXML element")
    supported_children = {_w("rPr"), _w("t")}
    for run in runs[first : last + 1]:
        unsupported = [child.tag for child in run if child.tag not in supported_children]
        if unsupported:
            raise WordSkillError(
                "tracked changes and comments refuse runs containing fields, tabs, "
                "breaks, drawings, or other non-text OOXML"
            )
    return runs, values, offsets, first, last


def _revision_range(
    paragraph: etree._Element,
    start: int,
    end: int,
    replacement: str,
    *,
    author: str,
    date: str,
    revision_id: int,
) -> int:
    runs, values, offsets, first, last = _affected_runs(paragraph, start, end)
    first_run = runs[first]
    insert_at = paragraph.index(first_run)
    prefix = values[first][: start - offsets[first]]
    suffix = values[last][end - offsets[last] :]

    deleted_parts: list[etree._Element] = []
    for index in range(first, last + 1):
        local_start = max(start - offsets[index], 0)
        local_end = min(end - offsets[index], len(values[index]))
        if local_end > local_start:
            deleted_parts.append(
                _text_run(runs[index], values[index][local_start:local_end], deleted=True)
            )

    for run in runs[first : last + 1]:
        paragraph.remove(run)

    additions: list[etree._Element] = []
    if prefix:
        additions.append(_text_run(first_run, prefix))
    deleted = etree.Element(
        _w("del"),
        {
            _w("id"): str(revision_id),
            _w("author"): author,
            _w("date"): date,
        },
    )
    deleted.extend(deleted_parts)
    additions.append(deleted)
    inserted = etree.Element(
        _w("ins"),
        {
            _w("id"): str(revision_id + 1),
            _w("author"): author,
            _w("date"): date,
        },
    )
    if replacement:
        inserted.append(_text_run(first_run, replacement))
    additions.append(inserted)
    if suffix:
        additions.append(_text_run(runs[last], suffix))
    for addition in additions:
        paragraph.insert(insert_at, addition)
        insert_at += 1
    return revision_id + 2


def _track_in_paragraph(
    paragraph: etree._Element,
    find: str,
    replacement: str,
    *,
    author: str,
    date: str,
    case_sensitive: bool,
    replace_all: bool,
    revision_id: int,
) -> tuple[int, int]:
    runs, values, _ = _direct_run_map(paragraph)
    full_text = "".join(values)
    starts = _matches(full_text, find, case_sensitive, replace_all)
    for start in reversed(starts):
        revision_id = _revision_range(
            paragraph,
            start,
            start + len(find),
            replacement,
            author=author,
            date=date,
            revision_id=revision_id,
        )
    return len(starts), revision_id


def _comment_range(
    paragraph: etree._Element, start: int, end: int, comment_id: int
) -> None:
    runs, values, offsets, first, last = _affected_runs(paragraph, start, end)
    first_run = runs[first]
    insert_at = paragraph.index(first_run)
    prefix = values[first][: start - offsets[first]]
    suffix = values[last][end - offsets[last] :]
    selected: list[etree._Element] = []
    for index in range(first, last + 1):
        local_start = max(start - offsets[index], 0)
        local_end = min(end - offsets[index], len(values[index]))
        if local_end > local_start:
            selected.append(_text_run(runs[index], values[index][local_start:local_end]))
    for run in runs[first : last + 1]:
        paragraph.remove(run)

    additions: list[etree._Element] = []
    if prefix:
        additions.append(_text_run(first_run, prefix))
    additions.append(etree.Element(_w("commentRangeStart"), {_w("id"): str(comment_id)}))
    additions.extend(selected)
    additions.append(etree.Element(_w("commentRangeEnd"), {_w("id"): str(comment_id)}))
    reference_run = etree.Element(_w("r"))
    properties = etree.SubElement(reference_run, _w("rPr"))
    etree.SubElement(properties, _w("rStyle"), {_w("val"): "CommentReference"})
    etree.SubElement(reference_run, _w("commentReference"), {_w("id"): str(comment_id)})
    additions.append(reference_run)
    if suffix:
        additions.append(_text_run(runs[last], suffix))
    for addition in additions:
        paragraph.insert(insert_at, addition)
        insert_at += 1


def _comment_in_paragraph(
    paragraph: etree._Element,
    find: str,
    *,
    case_sensitive: bool,
    replace_all: bool,
    next_comment_id: int,
) -> tuple[list[int], int]:
    _, values, _ = _direct_run_map(paragraph)
    full_text = "".join(values)
    starts = _matches(full_text, find, case_sensitive, replace_all)
    ids: list[int] = []
    for start in reversed(starts):
        comment_id = next_comment_id
        next_comment_id += 1
        _comment_range(paragraph, start, start + len(find), comment_id)
        ids.append(comment_id)
    ids.reverse()
    return ids, next_comment_id


def _max_word_id(entries: dict[str, bytes], local_name: str) -> int:
    ids: list[int] = []
    for name, data in entries.items():
        if not name.startswith("word/") or not name.endswith(".xml"):
            continue
        try:
            root = _xml(data, name)
        except WordSkillError:
            continue
        for element in root.xpath(f".//w:{local_name}", namespaces=NS):
            value = element.get(_w("id"))
            if value is not None and value.lstrip("-").isdigit():
                ids.append(int(value))
    return max(ids, default=-1)


def _comments_root(entries: dict[str, bytes]) -> etree._Element:
    data = entries.get("word/comments.xml")
    if data is not None:
        return _xml(data, "word/comments.xml")
    return etree.Element(_w("comments"), nsmap={"w": W_NS})


def _append_comment(
    root: etree._Element,
    comment_id: int,
    text: str,
    *,
    author: str,
    initials: str,
    date: str,
) -> None:
    comment = etree.SubElement(
        root,
        _w("comment"),
        {
            _w("id"): str(comment_id),
            _w("author"): author,
            _w("initials"): initials,
            _w("date"): date,
        },
    )
    paragraph = etree.SubElement(comment, _w("p"))
    paragraph_properties = etree.SubElement(paragraph, _w("pPr"))
    etree.SubElement(paragraph_properties, _w("pStyle"), {_w("val"): "CommentText"})
    annotation_run = etree.SubElement(paragraph, _w("r"))
    annotation_properties = etree.SubElement(annotation_run, _w("rPr"))
    etree.SubElement(
        annotation_properties, _w("rStyle"), {_w("val"): "CommentReference"}
    )
    etree.SubElement(annotation_run, _w("annotationRef"))
    text_run = etree.SubElement(paragraph, _w("r"))
    text_node = etree.SubElement(text_run, _w("t"))
    _set_text(text_node, text)


def _ensure_comments_package(entries: dict[str, bytes], comments: etree._Element) -> None:
    entries["word/comments.xml"] = _serialize(comments)
    relationships_name = "word/_rels/document.xml.rels"
    relationships = _xml(entries[relationships_name], relationships_name)
    existing = relationships.xpath(
        "./pr:Relationship[@Type=$type]",
        namespaces={"pr": REL_NS},
        type=COMMENTS_REL_TYPE,
    )
    if not existing:
        used_ids = {
            element.get("Id", "")
            for element in relationships.findall(_rel("Relationship"))
        }
        number = 1
        while f"rId{number}" in used_ids:
            number += 1
        etree.SubElement(
            relationships,
            _rel("Relationship"),
            {"Id": f"rId{number}", "Type": COMMENTS_REL_TYPE, "Target": "comments.xml"},
        )
        entries[relationships_name] = _serialize(relationships)

    content_types = _xml(entries["[Content_Types].xml"], "[Content_Types].xml")
    existing_override = content_types.xpath(
        "./ct:Override[@PartName='/word/comments.xml']",
        namespaces={"ct": CT_NS},
    )
    if not existing_override:
        etree.SubElement(
            content_types,
            _ct("Override"),
            {"PartName": "/word/comments.xml", "ContentType": COMMENTS_CONTENT_TYPE},
        )
        entries["[Content_Types].xml"] = _serialize(content_types)


def _spec(path: Path) -> dict[str, Any]:
    try:
        value = json.loads(path.read_text(encoding="utf-8"))
    except json.JSONDecodeError as exc:
        raise WordSkillError(f"invalid edit spec JSON: {exc}") from exc
    if not isinstance(value, dict):
        raise WordSkillError("edit spec must be a JSON object")
    schema = value.get("schema")
    if schema not in SUPPORTED_SCHEMAS:
        raise WordSkillError(
            "unknown edit spec schema; expected word-edit-spec/v1 or word-edit-spec/v1.1"
        )
    operations = value.get("operations")
    if not isinstance(operations, list) or not operations:
        raise WordSkillError("edit spec operations must be a non-empty array")
    allowed = SCHEMA_OPERATIONS[schema]
    for index, operation in enumerate(operations):
        if not isinstance(operation, dict) or operation.get("op") not in allowed:
            allowed_text = ", ".join(sorted(allowed))
            raise WordSkillError(
                f"operation {index} is not supported by {schema}; allowed: {allowed_text}"
            )
    return value


def _operation_strings(operation: dict[str, Any]) -> tuple[str, str]:
    if operation["op"] == "fill":
        find = operation.get("placeholder")
        replacement = operation.get("value")
    else:
        find = operation.get("find")
        replacement = operation.get("replace", "")
    if not isinstance(find, str) or not find:
        raise WordSkillError("each edit operation requires a non-empty target string")
    if not isinstance(replacement, str):
        raise WordSkillError("replacement/value must be a string")
    return find, replacement


def _write_package(entries: dict[str, bytes], output: Path) -> None:
    handle, temp_name = tempfile.mkstemp(
        prefix=f".{output.stem}-", suffix=".docx", dir=output.parent
    )
    os.close(handle)
    temp_path = Path(temp_name)
    try:
        with zipfile.ZipFile(temp_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
            for name, data in entries.items():
                archive.writestr(name, data)
        _atomic_replace(temp_path, output)
    finally:
        temp_path.unlink(missing_ok=True)


def edit_document(args: argparse.Namespace) -> dict[str, Any]:
    source = _input_docx(args.input)
    output = _output_docx(args.output, source=source, overwrite=args.overwrite)
    spec_path = _local_path(args.spec)
    spec = _spec(spec_path)
    entries = _load_package(source)
    required = {"word/document.xml", "word/_rels/document.xml.rels", "[Content_Types].xml"}
    missing = required.difference(entries)
    if missing:
        raise WordSkillError(f"DOCX is missing required parts: {', '.join(sorted(missing))}")

    roots: dict[str, etree._Element] = {}

    def root_for(name: str) -> etree._Element:
        if name not in roots:
            roots[name] = _xml(entries[name], name)
        return roots[name]

    revision_id = _max_word_id(entries, "ins") + 1
    revision_id = max(revision_id, _max_word_id(entries, "del") + 1)
    next_comment_id = _max_word_id(entries, "comment") + 1
    comments = _comments_root(entries)
    operation_results: list[dict[str, Any]] = []
    comments_added = False

    for operation_index, operation in enumerate(spec["operations"]):
        op = operation["op"]
        find, replacement = _operation_strings(operation)
        scope = operation.get("scope", "document")
        if op == "comment" and scope != "document":
            raise WordSkillError("comment currently supports only scope='document'")
        parts = _scoped_parts(entries, scope)
        if not parts:
            raise WordSkillError(f"scope {scope!r} did not select any DOCX parts")
        case_sensitive = operation.get("case_sensitive", True)
        replace_all = operation.get("all", True)
        if not isinstance(case_sensitive, bool) or not isinstance(replace_all, bool):
            raise WordSkillError("case_sensitive and all must be booleans")
        count = 0
        stop = False
        for part in parts:
            root = root_for(part)
            for paragraph in root.xpath(".//w:p", namespaces=NS):
                remaining_all = replace_all
                if op in {"replace", "fill"}:
                    changed = _replace_in_paragraph(
                        paragraph,
                        find,
                        replacement,
                        case_sensitive=case_sensitive,
                        replace_all=remaining_all,
                    )
                elif op == "track_replace":
                    author = operation.get("author", "Aeloon")
                    date = operation.get("date", _utc_now())
                    if not isinstance(author, str) or not author:
                        raise WordSkillError("track_replace author must be a non-empty string")
                    if not isinstance(date, str) or not date:
                        raise WordSkillError("track_replace date must be a non-empty string")
                    changed, revision_id = _track_in_paragraph(
                        paragraph,
                        find,
                        replacement,
                        author=author,
                        date=date,
                        case_sensitive=case_sensitive,
                        replace_all=remaining_all,
                        revision_id=revision_id,
                    )
                else:
                    author = operation.get("author", "Aeloon")
                    initials = operation.get("initials", "AE")
                    date = operation.get("date", _utc_now())
                    text = operation.get("text")
                    comment_values = (author, initials, date, text)
                    if not all(
                        isinstance(value, str) and value for value in comment_values
                    ):
                        raise WordSkillError(
                            "comment requires non-empty text; author, initials, and date "
                            "must be strings"
                        )
                    ids, next_comment_id = _comment_in_paragraph(
                        paragraph,
                        find,
                        case_sensitive=case_sensitive,
                        replace_all=remaining_all,
                        next_comment_id=next_comment_id,
                    )
                    for comment_id in ids:
                        _append_comment(
                            comments,
                            comment_id,
                            text,
                            author=author,
                            initials=initials,
                            date=date,
                        )
                    changed = len(ids)
                    comments_added = comments_added or bool(ids)
                count += changed
                if changed and not replace_all:
                    stop = True
                    break
            if stop:
                break
        if count == 0:
            raise WordSkillError(
                f"operation {operation_index} ({op}) found no occurrence of {find!r}"
            )
        operation_results.append({"index": operation_index, "op": op, "matches": count})

    for name, root in roots.items():
        entries[name] = _serialize(root)
    if comments_added:
        _ensure_comments_package(entries, comments)
    _write_package(entries, output)
    return {"status": "edited", "output": str(output), "operations": operation_results}


def _relationship_target(rel_name: str, target: str) -> str:
    if target.startswith("/"):
        return target.lstrip("/")
    path = PurePosixPath(rel_name)
    source_dir = PurePosixPath(".") if rel_name == "_rels/.rels" else path.parent.parent
    return posixpath.normpath(str(source_dir / target).split("#", 1)[0])


def validate_document(path: Path, *, readback: bool = True) -> dict[str, Any]:
    errors: list[str] = []
    warnings: list[str] = []
    try:
        entries = _load_package(path)
    except WordSkillError as exc:
        return {"valid": False, "path": str(path), "errors": [str(exc)], "warnings": []}

    required = {"[Content_Types].xml", "_rels/.rels", "word/document.xml"}
    for name in sorted(required.difference(entries)):
        errors.append(f"missing required OPC part: {name}")

    parsed: dict[str, etree._Element] = {}
    for name, data in entries.items():
        if name.endswith(".xml") or name.endswith(".rels"):
            try:
                parsed[name] = _xml(data, name)
            except WordSkillError as exc:
                errors.append(str(exc))

    for name, root in parsed.items():
        if not name.endswith(".rels"):
            continue
        for relationship in root.findall(_rel("Relationship")):
            if relationship.get("TargetMode") == "External":
                continue
            target = relationship.get("Target")
            if not target:
                errors.append(f"relationship without Target in {name}")
                continue
            resolved = _relationship_target(name, target)
            if resolved not in entries:
                errors.append(f"broken relationship in {name}: {target} -> {resolved}")

    document = parsed.get("word/document.xml")
    counts = {
        "paragraphs": 0,
        "tables": 0,
        "images": 0,
        "insertions": 0,
        "deletions": 0,
        "comments": 0,
    }
    extracted_text = ""
    if document is not None:
        counts["paragraphs"] = len(document.xpath(".//w:p", namespaces=NS))
        counts["tables"] = len(document.xpath(".//w:tbl", namespaces=NS))
        counts["images"] = len(document.xpath(".//w:drawing", namespaces=NS))
        counts["insertions"] = len(document.xpath(".//w:ins", namespaces=NS))
        counts["deletions"] = len(document.xpath(".//w:del", namespaces=NS))
        extracted_text = "".join(
            document.xpath(".//w:t[not(ancestor::w:del)]/text()", namespaces=NS)
        )
        for paragraph in document.xpath(".//w:p", namespaces=NS):
            properties = paragraph.find(_w("pPr"))
            if properties is not None and paragraph.index(properties) != 0:
                errors.append("w:pPr is not the first child of a paragraph")
                break
        for run in document.xpath(".//w:r", namespaces=NS):
            properties = run.find(_w("rPr"))
            if properties is not None and run.index(properties) != 0:
                errors.append("w:rPr is not the first child of a run")
                break
        body = document.find(_w("body"))
        if body is not None:
            section = body.find(_w("sectPr"))
            if section is not None and body.index(section) != len(body) - 1:
                errors.append("w:sectPr is not the final child of w:body")

    comment_root = parsed.get("word/comments.xml")
    comment_ids: set[str] = set()
    if comment_root is not None:
        comment_ids = {
            element.get(_w("id"), "")
            for element in comment_root.xpath("./w:comment", namespaces=NS)
        }
        counts["comments"] = len(comment_ids)
    if document is not None:
        reference_ids = {
            element.get(_w("id"), "")
            for element in document.xpath(".//w:commentReference", namespaces=NS)
        }
        range_start_ids = {
            element.get(_w("id"), "")
            for element in document.xpath(".//w:commentRangeStart", namespaces=NS)
        }
        range_end_ids = {
            element.get(_w("id"), "")
            for element in document.xpath(".//w:commentRangeEnd", namespaces=NS)
        }
        mismatched_comments = (
            reference_ids != comment_ids
            or range_start_ids != reference_ids
            or range_end_ids != reference_ids
        )
        if mismatched_comments:
            errors.append("comment definitions, ranges, and references do not have matching IDs")

    if "[待确认" in extracted_text:
        warnings.append("document contains unresolved [待确认：…] placeholders")

    semantic: dict[str, Any] = {"python_docx": "not_checked", "markitdown": "not_checked"}
    try:
        from docx import Document

        opened = Document(path)
        semantic["python_docx"] = "opened"
        semantic["python_docx_paragraphs"] = len(opened.paragraphs)
        semantic["python_docx_tables"] = len(opened.tables)
    except Exception as exc:  # package-specific parser errors vary
        errors.append(f"python-docx could not open the package: {exc}")
        semantic["python_docx"] = "failed"

    if readback:
        try:
            from markitdown import MarkItDown

            result = MarkItDown().convert_local(path)
            content = getattr(result, "text_content", None)
            if content is None:
                content = getattr(result, "markdown", "")
            semantic["markitdown"] = "read"
            semantic["markitdown_characters"] = len(content or "")
            if extracted_text.strip() and not (content or "").strip():
                errors.append("MarkItDown readback was empty for a non-empty document")
        except ImportError:
            warnings.append("MarkItDown is unavailable; content readback was skipped")
            semantic["markitdown"] = "unavailable"
        except Exception as exc:  # conversion backends expose heterogeneous errors
            warnings.append(f"MarkItDown readback failed: {exc}")
            semantic["markitdown"] = "failed"

    return {
        "valid": not errors,
        "path": str(path),
        "errors": errors,
        "warnings": warnings,
        "counts": counts,
        "semantic": semantic,
    }


def _write_json(path: str, value: dict[str, Any], *, overwrite: bool = False) -> Path:
    output = _local_path(path, must_exist=False)
    if output.exists() and not overwrite:
        raise WordSkillError(f"report already exists (pass --overwrite to replace it): {output}")
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(json.dumps(value, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    return output


def validate_action(args: argparse.Namespace) -> tuple[dict[str, Any], int]:
    source = _input_docx(args.input)
    report = validate_document(source, readback=not args.skip_readback)
    if args.report:
        _write_json(args.report, report, overwrite=args.overwrite)
    return report, 0 if report["valid"] else 2


def find_libreoffice() -> str | None:
    return shutil.which("soffice") or shutil.which("libreoffice")


def libreoffice_install_hint() -> str:
    if platform.system() == "Darwin":
        return "install LibreOffice with: brew install --cask libreoffice"
    if platform.system() == "Linux":
        return "install LibreOffice with: sudo apt-get install libreoffice"
    return "install LibreOffice and ensure soffice is available on PATH"


def render_document(args: argparse.Namespace) -> tuple[dict[str, Any], int]:
    executable = find_libreoffice()
    if not executable:
        report = {
            "status": "missing",
            "libreoffice": None,
            "visual_qa_complete": False,
            "hint": libreoffice_install_hint(),
        }
        return report, 2
    if args.check:
        return {
            "status": "available",
            "libreoffice": executable,
            "visual_qa_complete": False,
        }, 0
    if not args.input or not args.output_dir:
        raise WordSkillError("render requires INPUT and --output-dir unless --check is used")

    source = _input_docx(args.input)
    output_dir = _local_path(args.output_dir, must_exist=False)
    output_dir.mkdir(parents=True, exist_ok=True)
    existing_pages = list(output_dir.glob("page-*.png"))
    if existing_pages and not args.overwrite:
        raise WordSkillError(
            "render output directory contains page PNGs; pass --overwrite to replace them"
        )
    if args.overwrite:
        for page in existing_pages:
            page.unlink()

    try:
        import pypdfium2 as pdfium
    except ImportError as exc:
        raise WordSkillError("pypdfium2 is required to render DOCX page previews") from exc

    with tempfile.TemporaryDirectory(prefix="aeloon-word-render-") as temp_value:
        temp = Path(temp_value)
        profile = temp / "profile"
        converted = temp / f"{source.stem}.pdf"
        command = [
            executable,
            "--headless",
            f"-env:UserInstallation={profile.resolve().as_uri()}",
            "--convert-to",
            "pdf",
            "--outdir",
            str(temp),
            str(source),
        ]
        result = subprocess.run(command, capture_output=True, text=True, timeout=180)
        if result.returncode != 0 or not converted.is_file():
            detail = (result.stderr or result.stdout).strip()
            raise WordSkillError(f"LibreOffice DOCX conversion failed: {detail}")
        pdf = pdfium.PdfDocument(converted)
        pages: list[str] = []
        scale = args.dpi / 72
        try:
            for index in range(len(pdf)):
                page = pdf[index]
                try:
                    bitmap = page.render(scale=scale)
                    image = bitmap.to_pil()
                    output = output_dir / f"page-{index + 1:03d}.png"
                    image.save(output)
                    pages.append(str(output))
                finally:
                    page.close()
        finally:
            pdf.close()

    return {
        "status": "rendered",
        "input": str(source),
        "output_dir": str(output_dir),
        "pages": pages,
        "visual_qa_complete": False,
        "note": "Inspect every page PNG before claiming visual QA is complete.",
    }, 0


def parser() -> argparse.ArgumentParser:
    root = argparse.ArgumentParser(description=__doc__)
    actions = root.add_subparsers(dest="action", required=True)

    build = actions.add_parser("build", help="build DOCX from Markdown")
    build.add_argument("input")
    build.add_argument("output")
    build.add_argument("--title")
    build.add_argument("--subtitle")
    build.add_argument("--author")
    build.add_argument("--date")
    build.add_argument("--header")
    build.add_argument("--footer")
    build.add_argument("--toc", action="store_true")
    build.add_argument("--page-numbers", action="store_true")
    build.add_argument("--font", default="Arial")
    build.add_argument("--font-size", type=float, default=11.0)
    build.add_argument("--overwrite", action="store_true")

    edit = actions.add_parser("edit", help="edit DOCX using a versioned JSON spec")
    edit.add_argument("input")
    edit.add_argument("output")
    edit.add_argument("--spec", required=True)
    edit.add_argument("--overwrite", action="store_true")

    validate = actions.add_parser("validate", help="validate a DOCX package and content")
    validate.add_argument("input")
    validate.add_argument("--report")
    validate.add_argument("--skip-readback", action="store_true")
    validate.add_argument("--overwrite", action="store_true")

    render = actions.add_parser("render", help="render DOCX pages using LibreOffice")
    render.add_argument("input", nargs="?")
    render.add_argument("--output-dir")
    render.add_argument("--dpi", type=int, default=144)
    render.add_argument("--check", action="store_true")
    render.add_argument("--overwrite", action="store_true")
    return root


def main(argv: Sequence[str] | None = None) -> int:
    args = parser().parse_args(argv)
    try:
        if args.action == "build":
            report = build_document(args)
            code = 0
        elif args.action == "edit":
            report = edit_document(args)
            code = 0
        elif args.action == "validate":
            report, code = validate_action(args)
        else:
            report, code = render_document(args)
    except (WordSkillError, OSError, subprocess.SubprocessError) as exc:
        print(f"error: {exc}", file=sys.stderr)
        return 2
    print(json.dumps(report, ensure_ascii=False, indent=2))
    if args.action == "render" and code and report.get("hint"):
        print(report["hint"], file=sys.stderr)
    return code


if __name__ == "__main__":
    raise SystemExit(main())
