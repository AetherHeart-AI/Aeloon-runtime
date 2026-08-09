from __future__ import annotations

import copy
import importlib.util
import json
import sys
import zipfile
from argparse import Namespace
from pathlib import Path

import pytest
from docx import Document
from lxml import etree

RESOURCE_ROOT = (
    Path(__file__).parents[1] / "aeloon_core" / "resources" / "skills" / "word-docx"
)
CLI_PATH = RESOURCE_ROOT / "scripts" / "cli.py"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
NS = {"w": W_NS, "pr": REL_NS, "ct": CT_NS}


def load_cli():
    spec = importlib.util.spec_from_file_location("test_word_docx_cli", CLI_PATH)
    assert spec is not None and spec.loader is not None
    module = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = module
    try:
        spec.loader.exec_module(module)
    finally:
        sys.modules.pop(spec.name, None)
    return module


def save_source(path: Path, runs: list[tuple[str, bool]]) -> None:
    document = Document()
    paragraph = document.add_paragraph()
    for text, bold in runs:
        run = paragraph.add_run(text)
        run.bold = bold
    document.save(path)


def write_spec(path: Path, operations: list[dict[str, object]], version: str = "v1") -> None:
    path.write_text(
        json.dumps(
            {"schema": f"word-edit-spec/{version}", "operations": operations},
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )


def edit_args(source: Path, output: Path, spec: Path) -> Namespace:
    return Namespace(input=str(source), output=str(output), spec=str(spec), overwrite=False)


def test_skill_metadata_and_single_cli_entry_are_complete() -> None:
    skill = (RESOURCE_ROOT / "SKILL.md").read_text(encoding="utf-8")
    assert "present_files" in skill
    assert "word-edit-spec/v1.1" in skill
    assert "未读取、复制或改编" in skill
    assert (RESOURCE_ROOT / "LICENSE.txt").is_file()
    assert [path.name for path in (RESOURCE_ROOT / "scripts").glob("*.py")] == [
        "cli.py"
    ]


def test_build_markdown_creates_editable_docx(tmp_path: Path) -> None:
    cli = load_cli()
    source = tmp_path / "draft.md"
    source.write_text(
        "# 执行摘要\n\n这是 **重点** 和 [链接](https://example.com)。\n\n"
        "- 第一项\n- 第二项\n\n"
        "| 指标 | 值 |\n| --- | --- |\n| 收入 | 42 |\n\n"
        "<!-- pagebreak -->\n\n## 结论\n\n完成。\n",
        encoding="utf-8",
    )
    output = tmp_path / "report.docx"
    args = Namespace(
        input=str(source),
        output=str(output),
        title="季度报告",
        subtitle="内部材料",
        author="Aeloon",
        date="2026-08-08",
        header="机密",
        footer="内部使用",
        toc=True,
        page_numbers=True,
        font="Arial",
        font_size=11.0,
        overwrite=False,
    )

    report = cli.build_document(args)

    assert report["status"] == "created"
    opened = Document(output)
    assert any(paragraph.text == "季度报告" for paragraph in opened.paragraphs)
    assert any(paragraph.text == "执行摘要" for paragraph in opened.paragraphs)
    assert opened.tables[0].cell(1, 0).text == "收入"
    with zipfile.ZipFile(output) as archive:
        document_xml = archive.read("word/document.xml")
        footer_xml = archive.read("word/footer1.xml")
        relationships = archive.read("word/_rels/document.xml.rels")
    assert b"TOC" in document_xml
    assert b"PAGE" in footer_xml
    assert b"relationships/hyperlink" in relationships


def test_edit_replaces_text_split_across_runs_and_preserves_style(tmp_path: Path) -> None:
    cli = load_cli()
    source = tmp_path / "input.docx"
    output = tmp_path / "output.docx"
    spec = tmp_path / "edits.json"
    save_source(source, [("Hello ", True), ("world", False)])
    original = source.read_bytes()
    write_spec(
        spec,
        [{"op": "replace", "find": "lo world", "replace": "greetings"}],
    )

    report = cli.edit_document(edit_args(source, output, spec))

    assert report["operations"][0]["matches"] == 1
    assert source.read_bytes() == original
    opened = Document(output)
    assert opened.paragraphs[0].text == "Helgreetings"
    assert opened.paragraphs[0].runs[0].bold is True


def test_edit_writes_native_track_changes_across_runs(tmp_path: Path) -> None:
    cli = load_cli()
    source = tmp_path / "input.docx"
    output = tmp_path / "tracked.docx"
    spec = tmp_path / "tracked.json"
    save_source(source, [("The old ", True), ("value remains.", False)])
    write_spec(
        spec,
        [
            {
                "op": "track_replace",
                "find": "old value",
                "replace": "new value",
                "author": "Reviewer",
                "date": "2026-08-08T08:00:00Z",
                "all": False,
            }
        ],
        version="v1.1",
    )

    cli.edit_document(edit_args(source, output, spec))

    with zipfile.ZipFile(output) as archive:
        root = etree.fromstring(archive.read("word/document.xml"))
    deleted = root.xpath(".//w:del", namespaces=NS)
    inserted = root.xpath(".//w:ins", namespaces=NS)
    assert len(deleted) == len(inserted) == 1
    assert "".join(deleted[0].xpath(".//w:delText/text()", namespaces=NS)) == "old value"
    assert "".join(inserted[0].xpath(".//w:t/text()", namespaces=NS)) == "new value"
    assert deleted[0].get(f"{{{W_NS}}}author") == "Reviewer"
    assert deleted[0].xpath("./w:r[1]/w:rPr/w:b", namespaces=NS)

    # A Word-compatible consumer must be able to open the tracked package before
    # the user chooses whether to accept or reject the revision.
    Document(output)

    accepted = copy.deepcopy(root)
    for revision in accepted.xpath(".//w:del", namespaces=NS):
        revision.getparent().remove(revision)
    for revision in accepted.xpath(".//w:ins", namespaces=NS):
        parent = revision.getparent()
        position = parent.index(revision)
        for child in list(revision):
            parent.insert(position, child)
            position += 1
        parent.remove(revision)
    assert "".join(accepted.xpath(".//w:t/text()", namespaces=NS)) == (
        "The new value remains."
    )

    rejected = copy.deepcopy(root)
    for revision in rejected.xpath(".//w:ins", namespaces=NS):
        revision.getparent().remove(revision)
    for revision in rejected.xpath(".//w:del", namespaces=NS):
        parent = revision.getparent()
        position = parent.index(revision)
        for child in list(revision):
            for text_node in child.xpath(".//w:delText", namespaces=NS):
                text_node.tag = f"{{{W_NS}}}t"
            parent.insert(position, child)
            position += 1
        parent.remove(revision)
    assert "".join(rejected.xpath(".//w:t/text()", namespaces=NS)) == (
        "The old value remains."
    )


def test_edit_adds_comment_part_relationship_and_range(tmp_path: Path) -> None:
    cli = load_cli()
    source = tmp_path / "input.docx"
    output = tmp_path / "commented.docx"
    spec = tmp_path / "comments.json"
    save_source(source, [("Please verify this source.", False)])
    write_spec(
        spec,
        [
            {
                "op": "comment",
                "find": "verify this",
                "text": "请核对来源。",
                "author": "审阅人",
                "initials": "SR",
                "all": False,
            }
        ],
        version="v1.1",
    )

    cli.edit_document(edit_args(source, output, spec))

    with zipfile.ZipFile(output) as archive:
        names = set(archive.namelist())
        document = etree.fromstring(archive.read("word/document.xml"))
        comments = etree.fromstring(archive.read("word/comments.xml"))
        relationships = etree.fromstring(
            archive.read("word/_rels/document.xml.rels")
        )
        content_types = etree.fromstring(archive.read("[Content_Types].xml"))
    assert "word/comments.xml" in names
    assert document.xpath(".//w:commentRangeStart[@w:id='0']", namespaces=NS)
    assert document.xpath(".//w:commentRangeEnd[@w:id='0']", namespaces=NS)
    assert document.xpath(".//w:commentReference[@w:id='0']", namespaces=NS)
    assert comments.xpath("./w:comment[@w:id='0'][@w:author='审阅人']", namespaces=NS)
    assert "".join(comments.xpath(".//w:t/text()", namespaces=NS)) == "请核对来源。"
    assert relationships.xpath(
        "./pr:Relationship[contains(@Type, '/comments')]", namespaces=NS
    )
    assert content_types.xpath(
        "./ct:Override[@PartName='/word/comments.xml']", namespaces=NS
    )
    Document(output)


def test_validate_checks_structure_and_comment_integrity(tmp_path: Path) -> None:
    cli = load_cli()
    source = tmp_path / "valid.docx"
    save_source(source, [("A valid document", False)])

    report = cli.validate_document(source, readback=False)

    assert report["valid"] is True
    assert report["counts"]["paragraphs"] >= 1
    assert report["semantic"]["python_docx"] == "opened"


def test_edit_refuses_to_overwrite_input_or_miss_target(tmp_path: Path) -> None:
    cli = load_cli()
    source = tmp_path / "input.docx"
    spec = tmp_path / "edits.json"
    save_source(source, [("Known text", False)])
    write_spec(spec, [{"op": "replace", "find": "missing", "replace": "x"}])

    with pytest.raises(cli.WordSkillError, match="different"):
        cli.edit_document(edit_args(source, source, spec))
    with pytest.raises(cli.WordSkillError, match="found no occurrence"):
        cli.edit_document(edit_args(source, tmp_path / "output.docx", spec))


def test_render_missing_libreoffice_returns_install_hint(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    cli = load_cli()
    monkeypatch.setattr(cli, "find_libreoffice", lambda: None)
    monkeypatch.setattr(cli.platform, "system", lambda: "Darwin")

    report, code = cli.render_document(
        Namespace(
            input=None,
            output_dir=None,
            dpi=144,
            check=True,
            overwrite=False,
        )
    )

    assert code == 2
    assert report["visual_qa_complete"] is False
    assert report["hint"] == "install LibreOffice with: brew install --cask libreoffice"
